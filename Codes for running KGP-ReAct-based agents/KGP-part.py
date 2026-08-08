import os
import json
import re
from typing import List

import pandas as pd
import torch
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from langchain.schema import Document
from langchain_core.embeddings import Embeddings
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain


class QZhouEmbeddings(Embeddings):
    """LangChain-compatible wrapper for Kingsoft-LLM/QZhou-Embedding."""

    def __init__(
        self,
        model_name: str = "Kingsoft-LLM/QZhou-Embedding",
        batch_size: int = 1,
    ):
        self.batch_size = batch_size
        self.device = "cuda" if torch.cuda.is_available() else "cpu"

        model_kwargs = {
            "device_map": self.device,
            "trust_remote_code": True,
        }
        if self.device == "cuda":
            model_kwargs.update({
                "torch_dtype": (
                    torch.bfloat16
                    if torch.cuda.is_bf16_supported()
                    else torch.float16
                ),
                "attn_implementation": "sdpa",
            })

        self.model = SentenceTransformer(
            model_name,
            model_kwargs=model_kwargs,
            tokenizer_kwargs={
                "padding_side": "left",
                "trust_remote_code": True,
            },
            trust_remote_code=True,
        )

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Generate normalized embeddings for candidate-chain documents."""
        embeddings = self.model.encode(
            texts,
            batch_size=self.batch_size,
            normalize_embeddings=True,
            convert_to_numpy=True,
            show_progress_bar=False,
        )
        return embeddings.tolist()

    def embed_query(self, text: str) -> List[float]:
        """Generate a normalized query embedding using QZhou's query prompt."""
        embedding = self.model.encode(
            [text],
            prompt_name="query",
            batch_size=1,
            normalize_embeddings=True,
            convert_to_numpy=True,
            show_progress_bar=False,
        )[0]
        return embedding.tolist()


# —— 第一个：从 JSON 生成 Word chunks —— #
def tool_graph_to_chunks(input_json_file):
    """
    Processes the tool relationship data from 'tool_relationship.json' and generates
    a Word document ('tool_chunks.docx') containing the tool relationships, formatted as chains.
    """
    output_doc_file = r"C:\Users\26389\OneDrive\Desktop\tool\tool_chunks.docx"
    output_step_doc_file = r"C:\Users\26389\OneDrive\Desktop\tool\step_chunks.docx"

    with open(input_json_file, 'r', encoding='utf-8-sig') as f:
        data = json.load(f)

    # 在本地把 DocxDocument 赋值给 Document，以保持函数内部代码不变
    Document = DocxDocument
    doc = Document()

    def add_relationship_chain_to_doc(chain_str, chain_name):
        doc.add_paragraph(f"{chain_name}. {chain_str}")

    relationships = []
    step_relationships = []

    for chain in data:
        chain_name = chain["chainName"]  # Each chain has a name
        chain_steps = chain["chainSteps"]  # Steps in the chain

        # Create a list to hold the tool chain
        tool_chain = []
        step_chain = []

        # Collect all tools in the current chain
        for step in chain_steps:
            tool_chain.append(step["tool"])
            step_chain.append(step["step"])

        full_chain_str = ", and then ".join(tool_chain)
        full_step_str = ", ".join(step_chain)
        relationships.append((chain_name, full_chain_str))
        step_relationships.append((chain_name, full_step_str))

        # Output each full relationship chain with a number
    for chain_name, chain_str in relationships:
        add_relationship_chain_to_doc(chain_str, chain_name)

        # Save the generated document to the specified output file
    doc.save(output_doc_file)

    print(f"Document saved as {output_doc_file}")
    step_doc = Document()
    for chain_name, chain_str in step_relationships:
        step_doc.add_paragraph(f"{chain_name}. {chain_str}")
    step_doc.save(output_step_doc_file)
    print(f"Document saved as {output_step_doc_file}")



    return output_doc_file

# —— 第二个：RAG 检索 —— #
def rag(input_docx):
    # 原代码中的 Excel 路径与 API Key
    input_questions_excel = r"C:\Users\26389\OneDrive\Desktop\tool\Task.xlsx"
    os.environ["OPENAI_API_KEY"] = ""

    def split_tasks_by_regex(text):
        pattern = re.compile(r'(Task\d+\.)')
        parts = pattern.split(text)
        tasks = []
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                task_full = parts[i] + parts[i + 1].strip()
                tasks.append(task_full)
        return tasks

    def extract_text_from_docx(file_path):
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

    text = extract_text_from_docx(input_docx)
    # 通过 Task 分界切分文本
    task_texts = split_tasks_by_regex(text)

    # 构造包含 Task 信息的 Document 列表
    all_splits = []
    for task_text in task_texts:
        task_id_match = re.search(r'(Task\d+)\.', task_text)
        task_id = task_id_match.group(1) if task_id_match else "unknown"
        all_splits.append(Document(page_content=task_text, metadata={"task_id": task_id, "source": input_docx}))

    # Load the questions from the Excel file
    questions_df = pd.read_excel(input_questions_excel)

    # Set up QZhou-Embedding and vector store
    embedding_model = QZhouEmbeddings(
        model_name="Kingsoft-LLM/QZhou-Embedding"
    )


    vector_store = FAISS.from_documents(all_splits, embedding_model)

    # Prepare the output DataFrame
    output_df = pd.DataFrame(columns=["task","question", "retrieved_task_1", "retrieved_task_2", "retrieved_task_3", "retrieved_task_4", "retrieved_task_5", "retrieved_task_6", "retrieved_task_7", "task_exist"])

    # Retrieve chunks for each question
    for index, row in questions_df.iterrows():
        user_question = row["question"]
        task_text = row["task"]

        # Retrieve relevant documents using similarity search
        retrieved_docs = vector_store.similarity_search(user_question, k=7)

        # Extract the context from retrieved chunks
        task_exist = 'yes' if any(task_text in doc.page_content for doc in retrieved_docs) else 'no'


        # Prepare a new DataFrame for the new row
        new_row = pd.DataFrame({
            "task": [task_text],
            "question": [user_question],
            "retrieved_task_1": [retrieved_docs[0].page_content],
            "retrieved_task_2": [retrieved_docs[1].page_content],
            "retrieved_task_3": [retrieved_docs[2].page_content],
            "retrieved_task_4": [retrieved_docs[3].page_content],
            "retrieved_task_5": [retrieved_docs[4].page_content],
            "retrieved_task_6": [retrieved_docs[5].page_content],
            "retrieved_task_7": [retrieved_docs[6].page_content],
            "task_exist": [task_exist]
        })

        # Concatenate the new row to the output DataFrame
        output_df = pd.concat([output_df, new_row], ignore_index=True)

    # Save the output DataFrame to Excel
    output_file_path = r"C:\Users\26389\OneDrive\Desktop\tool\KGP_derived_candidate_chain.xlsx"
    output_df.to_excel(output_file_path, index=False)

    print(f"RAG results saved to {output_file_path}")
    return output_file_path

chunks_docx = tool_graph_to_chunks(input_json_file=r"C:/Users/26389/Desktop/neo4j_query_table_data.json")
excel_with_retrieved = rag(input_docx= r"C:\Users\26389\OneDrive\Desktop\tool\tool_chunks.docx")
# 读取包含问题和候选链的 Excel 文件
excel_path = r"C:\Users\26389\OneDrive\Desktop\tool\KGP_derived_candidate_chain.xlsx"
df = pd.read_excel(excel_path)

# 创建模型实例
os.environ["OPENAI_API_KEY"] = ""

llm = ChatOpenAI(
    openai_api_key=os.environ["OPENAI_API_KEY"],
    temperature=0,
    model_name="gpt-5"
)

# 定义提示模板
template = """
You are an intelligent assistant. Your task is to choose the most relevant candidate chain that best solves the user's question, based on the detailed information provided for each candidate chain.
– First, discard any chain that contains steps beyond what’s strictly needed to answer the user’s question.
– Among the remaining, choose the one whose steps exactly solve the question with no extras.
The user question is:
{question}

The following are the details of the candidate chains:
1. {retrieved_task_1}
2. {retrieved_task_2}
3. {retrieved_task_3}
4. {retrieved_task_4}
5. {retrieved_task_5}
6. {retrieved_task_6}
7. {retrieved_task_7}

Please focus on the information of each candidate chains when selecting the most relevant chain for the given question. Based on the details provided, choose the chain that best solves the user’s question.

Please return all the information of the candidate chain you selected.
"""

# 创建提示模板
prompt = PromptTemplate(
    input_variables=[
        "question",
        "retrieved_task_1",
        "retrieved_task_2",
        "retrieved_task_3",
        "retrieved_task_4",
        "retrieved_task_5",
        "retrieved_task_6",
        "retrieved_task_7"
    ],
    template=template,
)

# 将提示模板和语言模型组合成链式调用
llm_chain = LLMChain(prompt=prompt, llm=llm)

# 用于实时保存每个问题的回答结果
results = []

# 遍历 Excel 的每一行，处理每个问题
for index, row in df.iterrows():
    # 格式化提示内容
    candidate_chain_inputs = {
        "question": row["question"],
        "retrieved_task_1": row["retrieved_task_1"],
        "retrieved_task_2": row["retrieved_task_2"],
        "retrieved_task_3": row["retrieved_task_3"],
        "retrieved_task_4": row["retrieved_task_4"],
        "retrieved_task_5": row["retrieved_task_5"],
        "retrieved_task_6": row["retrieved_task_6"],
        "retrieved_task_7": row["retrieved_task_7"]
    }

    # 调用 llm_chain 生成结果
    response = llm_chain.invoke(candidate_chain_inputs)

    # 保存结果到 results 列表中
    results.append({
        "question": row["question"],
        "response": response["text"]  # 假设返回的结果是以 'text' 字段为键
    })



    # 将结果实时写入 Excel
    results_df = pd.DataFrame(results)
    results_df.to_excel(r"C:\Users\26389\OneDrive\Desktop\tool\LLM_choice_results.xlsx", index=False)

print("The LLM selection results have been saved to the file")

# 读取 Excel 文件
file_path = r"C:\Users\26389\OneDrive\Desktop\tool\LLM_choice_results.xlsx"
df = pd.read_excel(file_path)


def extract_matched_word(response):
    # 用 \b 保证 Task\d+ 是独立单词
    matches = re.findall(r'\bTask\d+\b', str(response))
    # 保留出现顺序的唯一值
    unique = list(dict.fromkeys(matches))
    if len(unique) == 0:
        return ''
    elif len(unique) == 1:
        return unique[0]
    else:
        return 'The responses involve multiple answers, require manual checking'

# 应用函数，生成新列 'matched_word'
df['matched_word'] = df['response'].apply(extract_matched_word)

# 保存回 Excel（不写入行索引）
df.to_excel(file_path, index=False)

# —— 一、读取 Excel 并提取 matched_word 列 —— #
excel_path = r"C:\Users\26389\OneDrive\Desktop\tool\LLM_choice_results.xlsx"
df = pd.read_excel(excel_path, dtype={'matched_word': str})
matched_words = df['matched_word'].tolist()
raw = df['matched_word'].dropna().astype(str).tolist()
matched_words = [w for w in raw if re.fullmatch(r'Task\d+', w)]
# —— 二、读取 Word 文档段落的通用函数 —— #
def load_paragraphs(docx_path):
    doc = DocxDocument(docx_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# 分别加载两个文档
paragraphs_0 = load_paragraphs(r"C:\Users\26389\OneDrive\Desktop\tool\step_chunks.docx")
paragraphs_1 = load_paragraphs(r"C:\Users\26389\OneDrive\Desktop\tool\tool_chunks.docx")

# —— 三、提取以 TaskN. 开头的整段 —— #
def extract_task_paragraphs(paragraphs):
    results = {}
    for word in matched_words:
        for para in paragraphs:
            if para.startswith(word + "."):
                results[word] = para
                break
    return results

task_results_0 = extract_task_paragraphs(paragraphs_0)
task_results_1 = extract_task_paragraphs(paragraphs_1)

df['Result_step'] = df['matched_word'].map(task_results_0)
df['Result_1']    = df['matched_word'].map(task_results_1)

# —— 四、格式化各自结果 —— #

# （1）第一个文档格式化：保留之前的英文序号逻辑
def format_with_ordinal(result: str, delimiter: str):
    if not isinstance(result, str) or not result:
        return ""
    content = re.sub(r"^Task\d+\.\s*", "", result)
    parts = content.split(delimiter)
    formatted = []
    for i, part in enumerate(parts, start=1):
        # 英文序数
        if i <= 10:
            names = ['First','Second','Third','Fourth','Fifth',
                     'Sixth','Seventh','Eighth','Ninth','Tenth']
            ordinal = names[i-1]
        else:
            suffix = ('th' if 11 <= (i % 100) <= 13 else
                      {1:'st',2:'nd',3:'rd'}.get(i % 10, 'th'))
            ordinal = f"{i}{suffix}"
        formatted.append(f"({i}) {ordinal}, {part.strip()}")
    return "\n".join(formatted)

df['Results_step'] = df['Result_step'].apply(lambda r: format_with_ordinal(r, ', '))

# （2）第二个文档格式化：按 ", and then " 拆分，每条后加 "_tool"
def format_for_tool(result: str, delimiter: str):
    if not isinstance(result, str) or not result:
        return ""
    content = re.sub(r"^Task\d+\.\s*", "", result)
    parts = content.split(delimiter)
    # 给每条加后缀并换行
    return "\n".join(part.strip() + "_tool" for part in parts)

df['Results_1'] = df['Result_1'].apply(lambda r: format_for_tool(r, ', and then '))

# —— 五、保存到新的 Excel —— #
output_excel_path = r"C:\Users\26389\OneDrive\Desktop\tool\LLM_results.xlsx"
df.to_excel(output_excel_path, index=False)
print("The recommended steps and the toolset have been saved to the file")
